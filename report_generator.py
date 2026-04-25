"""
研究报告自动生成模块
支持生成 Word 文档和 Excel 表格
"""

import os
import re
from datetime import datetime
from typing import List, Dict, Optional
from collections import Counter

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import OUTPUT_DIR, DEFAULT_REPORT_NAME


class ReportGenerator:
    """研究报告生成器"""

    def __init__(self, output_dir: str = OUTPUT_DIR):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_excel(
        self,
        articles: List[Dict],
        filename: Optional[str] = None
    ) -> str:
        """
        生成Excel文献汇总表

        Args:
            articles: 文献列表
            filename: 输出文件名

        Returns:
            生成的文件路径
        """
        if not filename:
            filename = f"文献汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        filepath = os.path.join(self.output_dir, filename)

        # 准备数据
        data = []
        for idx, article in enumerate(articles, 1):
            data.append({
                "序号": idx,
                "PMID": article.get("pmid", ""),
                "标题": article.get("title", ""),
                "第一作者": article.get("first_author", ""),
                "作者列表": ", ".join(article.get("authors", [])),
                "期刊": article.get("journal", ""),
                "发表日期": article.get("pub_date", ""),
                "年份": article.get("year", ""),
                "DOI": article.get("doi", ""),
                "关键词": ", ".join(article.get("keywords", [])),
                "MeSH主题词": ", ".join(article.get("mesh_terms", [])),
                "发表类型": ", ".join(article.get("publication_types", [])),
                "PubMed链接": article.get("url", ""),
                "摘要": article.get("abstract", "")
            })

        df = pd.DataFrame(data)

        # 创建Excel writer
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='文献列表', index=False)

            # 获取工作表
            worksheet = writer.sheets['文献列表']

            # 调整列宽
            column_widths = {
                'A': 8,   # 序号
                'B': 12,  # PMID
                'C': 60,  # 标题
                'D': 20,  # 第一作者
                'E': 40,  # 作者列表
                'F': 30,  # 期刊
                'G': 15,  # 发表日期
                'H': 10,  # 年份
                'I': 25,  # DOI
                'J': 40,  # 关键词
                'K': 40,  # MeSH主题词
                'L': 30,  # 发表类型
                'M': 30,  # PubMed链接
                'N': 80   # 摘要
            }

            for col, width in column_widths.items():
                worksheet.column_dimensions[col].width = width

        print(f"Excel文献汇总已生成: {filepath}")
        return filepath

    def generate_word_report(
        self,
        articles: List[Dict],
        query: str,
        filename: Optional[str] = None,
        title: Optional[str] = None
    ) -> str:
        """
        生成Word研究报告

        Args:
            articles: 文献列表
            query: 检索词
            filename: 输出文件名
            title: 报告标题

        Returns:
            生成的文件路径
        """
        if not filename:
            filename = f"研究报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        filepath = os.path.join(self.output_dir, filename)

        if not title:
            title = f"{DEFAULT_REPORT_NAME} - {query}"

        doc = Document()

        # 设置默认字体
        self._set_document_font(doc)

        # 1. 封面/标题
        self._add_title(doc, title)

        # 2. 检索信息
        self._add_search_info(doc, query, len(articles))

        # 3. 统计分析
        self._add_statistics(doc, articles)

        # 4. 文献列表
        self._add_article_list(doc, articles)

        # 5. 详细摘要
        self._add_abstracts(doc, articles)

        doc.save(filepath)
        print(f"Word研究报告已生成: {filepath}")
        return filepath

    def _set_document_font(self, doc: Document):
        """设置文档默认字体"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_title(self, doc: Document, title: str):
        """添加报告标题"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)

        # 添加生成日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # 空行

    def _add_search_info(self, doc: Document, query: str, total: int):
        """添加检索信息"""
        doc.add_heading("一、检索信息", level=1)

        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表头
        cells = info_table.rows[0].cells
        cells[0].text = "检索词"
        cells[1].text = "检索结果数量"

        cells = info_table.rows[1].cells
        cells[0].text = query
        cells[1].text = f"{total} 篇"

        # 设置表格样式
        for row in info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        doc.add_paragraph()

    def _add_statistics(self, doc: Document, articles: List[Dict]):
        """添加统计分析"""
        doc.add_heading("二、统计分析", level=1)

        if not articles:
            doc.add_paragraph("无文献数据")
            return

        # 1. 年份分布
        years = [a.get("year", "") for a in articles if a.get("year")]
        if years:
            doc.add_heading("2.1 发表年份分布", level=2)
            year_counts = Counter(years)
            year_table = doc.add_table(rows=1, cols=2)
            year_table.style = 'Table Grid'

            hdr_cells = year_table.rows[0].cells
            hdr_cells[0].text = "年份"
            hdr_cells[1].text = "文献数量"

            for year, count in sorted(year_counts.items(), reverse=True):
                row_cells = year_table.add_row().cells
                row_cells[0].text = year
                row_cells[1].text = str(count)

            doc.add_paragraph()

        # 2. 期刊分布
        journals = [a.get("journal", "") for a in articles if a.get("journal")]
        if journals:
            doc.add_heading("2.2 主要来源期刊", level=2)
            journal_counts = Counter(journals)

            journal_table = doc.add_table(rows=1, cols=3)
            journal_table.style = 'Table Grid'

            hdr_cells = journal_table.rows[0].cells
            hdr_cells[0].text = "排名"
            hdr_cells[1].text = "期刊名称"
            hdr_cells[2].text = "文献数量"

            for idx, (journal, count) in enumerate(journal_counts.most_common(10), 1):
                row_cells = journal_table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = journal
                row_cells[2].text = str(count)

            doc.add_paragraph()

        # 3. 作者统计
        all_authors = []
        for article in articles:
            all_authors.extend(article.get("authors", []))

        if all_authors:
            doc.add_heading("2.3 主要作者", level=2)
            author_counts = Counter(all_authors)

            author_table = doc.add_table(rows=1, cols=3)
            author_table.style = 'Table Grid'

            hdr_cells = author_table.rows[0].cells
            hdr_cells[0].text = "排名"
            hdr_cells[1].text = "作者"
            hdr_cells[2].text = "发文数量"

            for idx, (author, count) in enumerate(author_counts.most_common(10), 1):
                row_cells = author_table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = author
                row_cells[2].text = str(count)

            doc.add_paragraph()

        # 4. 关键词分析
        all_keywords = []
        for article in articles:
            all_keywords.extend(article.get("keywords", []))
            all_keywords.extend(article.get("mesh_terms", []))

        if all_keywords:
            doc.add_heading("2.4 高频关键词/主题词", level=2)
            keyword_counts = Counter(all_keywords)

            kw_table = doc.add_table(rows=1, cols=3)
            kw_table.style = 'Table Grid'

            hdr_cells = kw_table.rows[0].cells
            hdr_cells[0].text = "排名"
            hdr_cells[1].text = "关键词"
            hdr_cells[2].text = "出现频次"

            for idx, (kw, count) in enumerate(keyword_counts.most_common(15), 1):
                row_cells = kw_table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = kw
                row_cells[2].text = str(count)

            doc.add_paragraph()

    def _add_article_list(self, doc: Document, articles: List[Dict]):
        """添加文献列表"""
        doc.add_heading("三、文献列表", level=1)

        if not articles:
            doc.add_paragraph("无文献数据")
            return

        for idx, article in enumerate(articles, 1):
            # 标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{idx}. {article.get('title', '无标题')}")
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = RGBColor(0, 51, 102)

            # 作者
            authors = article.get("authors", [])
            if authors:
                author_text = ", ".join(authors[:5])
                if len(authors) > 5:
                    author_text += " 等"
                doc.add_paragraph(f"作者: {author_text}")

            # 期刊和日期
            journal = article.get("journal", "")
            pub_date = article.get("pub_date", "")
            if journal or pub_date:
                doc.add_paragraph(f"期刊: {journal} | 发表日期: {pub_date}")

            # PMID和链接
            pmid = article.get("pmid", "")
            if pmid:
                doc.add_paragraph(f"PMID: {pmid} | 链接: {article.get('url', '')}")

            # DOI
            doi = article.get("doi", "")
            if doi:
                doc.add_paragraph(f"DOI: {doi}")

            # 关键词
            keywords = article.get("keywords", [])
            mesh_terms = article.get("mesh_terms", [])
            all_terms = keywords + mesh_terms
            if all_terms:
                doc.add_paragraph(f"关键词: {', '.join(all_terms[:10])}")

            doc.add_paragraph()  # 文献间空行

    def _add_abstracts(self, doc: Document, articles: List[Dict]):
        """添加详细摘要"""
        doc.add_page_break()
        doc.add_heading("四、文献摘要", level=1)

        if not articles:
            doc.add_paragraph("无文献数据")
            return

        for idx, article in enumerate(articles, 1):
            # 标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{idx}. {article.get('title', '无标题')}")
            title_run.bold = True
            title_run.font.size = Pt(12)

            # 摘要
            abstract = article.get("abstract", "")
            if abstract:
                doc.add_heading("摘要:", level=3)
                # 清理摘要格式
                abstract = re.sub(r'\s+', ' ', abstract).strip()
                abstract_para = doc.add_paragraph(abstract)
                abstract_para.paragraph_format.line_spacing = 1.5
            else:
                doc.add_paragraph("[无摘要]")

            doc.add_paragraph()  # 空行

    def generate_both(
        self,
        articles: List[Dict],
        query: str,
        report_title: Optional[str] = None
    ) -> tuple:
        """
        同时生成Excel和Word报告

        Returns:
            (excel_path, word_path)
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        excel_path = self.generate_excel(
            articles,
            f"文献汇总_{timestamp}.xlsx"
        )

        word_path = self.generate_word_report(
            articles,
            query,
            f"研究报告_{timestamp}.docx",
            report_title
        )

        return excel_path, word_path


def generate_report(
    articles: List[Dict],
    query: str,
    output_format: str = "both",
    report_title: Optional[str] = None
) -> List[str]:
    """
    便捷的报告生成函数

    Args:
        articles: 文献列表
        query: 检索词
        output_format: 输出格式 (excel/word/both)
        report_title: 报告标题

    Returns:
        生成的文件路径列表
    """
    generator = ReportGenerator()
    paths = []

    if output_format in ["excel", "both"]:
        excel_path = generator.generate_excel(articles)
        paths.append(excel_path)

    if output_format in ["word", "both"]:
        word_path = generator.generate_word_report(
            articles, query, title=report_title
        )
        paths.append(word_path)

    return paths


if __name__ == "__main__":
    # 测试示例
    test_articles = [
        {
            "pmid": "12345678",
            "title": "Test Article Title",
            "abstract": "This is a test abstract for demonstration purposes.",
            "authors": ["Smith J", "Doe A"],
            "first_author": "Smith J",
            "journal": "Test Journal",
            "pub_date": "2023-01-15",
            "year": "2023",
            "doi": "10.1234/test.123",
            "keywords": ["test", "demo"],
            "mesh_terms": ["Testing"],
            "publication_types": ["Journal Article"],
            "url": "https://pubmed.ncbi.nlm.nih.gov/12345678/"
        }
    ]

    generate_report(test_articles, "test query", output_format="both")
